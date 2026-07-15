#!/usr/bin/env python3
"""Build LSRI narrative DOCX from canonical markdown (RFP typeset spec)."""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

HERE = Path(__file__).parent
MD = HERE / "lsri-narrative-word-paste-2026-07-10.md"
OUT = HERE / "lsri-application-narrative-2026-07-14.docx"


def normalize(text: str) -> str:
    return (
        text.replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
    )


def setup_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return doc


def add_run_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(normalize(text))
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_mixed_paragraph(doc: Document, text: str) -> None:
    """Handle **bold** segments and *italic* segments."""
    p = doc.add_paragraph()
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)")
    for part in pattern.findall(text):
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(normalize(part[2:-2]))
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(normalize(part[1:-1]))
            run.italic = True
        else:
            run = p.add_run(normalize(part))
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            table.rows[r].cells[c].text = normalize(cell.strip())
            for para in table.rows[r].cells[c].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)


def parse_md_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if not line.strip().startswith("|"):
            break
        if re.match(r"^\|\s*-+", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def build() -> Path:
    raw = MD.read_text(encoding="utf-8")
    # Strip HTML comments and fidelity footer before references
    raw = re.sub(r"<!--.*?-->", "", raw, flags=re.S)
    if "# References" in raw:
        body_text, refs_text = raw.split("# References", 1)
    else:
        body_text, refs_text = raw, ""

    # Remove trailing source fidelity block from body
    body_text = re.sub(
        r"\nSource entries used:.*",
        "",
        body_text,
        flags=re.S,
    ).strip()

    doc = setup_doc()
    lines = body_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(normalize(title))
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(normalize(line[3:].strip()), level=1)
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_md_table(table_lines))
            continue
        if line.startswith("**") and line.endswith("**") and ":" in line:
            add_mixed_paragraph(doc, line)
            i += 1
            continue
        add_mixed_paragraph(doc, line)
        i += 1

    # References on separate page (beyond 4-page limit)
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    doc.add_heading("References", level=1)
    for line in refs_text.splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        para = doc.add_paragraph()
        run = para.add_run(normalize(line))
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote: {path}")
