#!/usr/bin/env python3
"""Build LSRI Word documents: Yazdi one-pager + full application narrative."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUT_DIR = Path(__file__).parent


def normalize(text: str) -> str:
    return (
        text.replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
    )


def setup_doc(margins_in: float = 0.5) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(margins_in)
        section.bottom_margin = Inches(margins_in)
        section.left_margin = Inches(margins_in)
        section.right_margin = Inches(margins_in)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return doc


def add_title(doc: Document, text: str, size: int = 14) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(normalize(text))
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_subtitle(doc: Document, text: str, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(normalize(text))
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_heading(doc: Document, text: str, level: int = 2) -> None:
    h = doc.add_heading(normalize(text), level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"


def add_label_para(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(normalize(label))
    r1.bold = True
    r1.font.name = "Times New Roman"
    r2 = p.add_run(normalize(body))
    r2.font.name = "Times New Roman"


def add_para(doc: Document, text: str, bold: bool = False, size: int = 12) -> None:
    p = doc.add_paragraph()
    run = p.add_run(normalize(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(normalize(item), style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(normalize(item), style="List Number")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = normalize(h)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = normalize(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build_yazdi_one_pager() -> Path:
    doc = setup_doc(0.75)
    doc.styles["Normal"].font.size = Pt(11)

    add_title(
        doc,
        "Proposed CBID Work Package: AI-Supported Health Navigation for Refugee Populations",
        13,
    )
    add_subtitle(
        doc,
        "JHU Life Sciences Research Initiative, High-Impact Individual Award  |  "
        "PI: Shatha Elnakib (Center for Humanitarian Health, BSPH)  |  "
        "$200,000-$500,000 over 24 months  |  Submission: July 15, 2026",
        10,
    )
    doc.add_paragraph()

    add_heading(doc, "The study in brief", 2)
    add_para(
        doc,
        "We will test whether routinely generated, de-identified conversations from an operating "
        "digital health service can identify what blocks refugee and returnee women from using "
        "relevant public health services, and whether a navigation feature designed around one "
        "validated barrier is feasible and associated with changes in service use. The platform "
        "is HERA Digital Health's WhatsApp-based service for displaced populations; CHH leads "
        "design, ethics, analysis, and evaluation.",
        size=11,
    )
    add_para(
        doc,
        "Three aims over 24 months: (1) build and validate a governed conversational-and-utilization "
        "evidence pipeline under JHU IRB; (2) identify and prioritize modifiable barriers to service "
        "use; (3) design and feasibility-test one AI-supported navigation intervention on the live "
        "platform using rapid-cycle implementation evaluation.",
        size=11,
    )

    add_heading(doc, "The CBID role: one bounded design cycle", 2)
    add_para(
        doc,
        "CBID students would enter in the second year with validated findings in hand, not a cold start. "
        "The cycle: translate one CHH-approved barrier finding into user, accessibility, and system "
        "requirements; map the patient navigation journey; generate and assess concepts; produce one "
        "bounded HERA-compatible prototype or workflow; run supervised usability and feasibility testing "
        "under the approved ethics protocol; and deliver design documentation and an implementation handoff.",
        size=11,
    )
    add_para(
        doc,
        "The faculty role is design-methodology oversight, student supervision, translational feasibility "
        "advice, and milestone reviews, sized to your availability. The role could be scoped as named "
        "co-investigator or faculty advisor, subject to your preference and JHU eligibility requirements.",
        size=11,
    )

    add_heading(doc, "Why we think this fits CBID", 2)
    add_para(
        doc,
        "CBID fills a specific translational gap rather than duplicating the research team. Elnakib's "
        "published work on refugee integration into national health systems and humanitarian "
        "reproductive, maternal, and newborn health grounds the barrier framework and outcome "
        "interpretation. Surmeli's registered antenatal-care trial (NCT05094518) provides direct "
        "experience with digital navigation, utilization endpoints, and the data-linkage failures that "
        "can arise in routine refugee-clinic operations. CHH and the DrPH trainee therefore bring a "
        "defined research question and an evidence-governance plan; CBID would own the needs-driven "
        "translation of one validated finding into a usable workflow.",
        size=11,
    )
    add_para(
        doc,
        "This partner-anchored sequence is consistent with CBID's student design model and its published "
        "needs-finding approach for health-service innovation in low- and middle-income countries. "
        "The design constraints are substantive: multilingual and low-bandwidth use, privacy, accessibility, "
        "bias awareness, and protection-sensitive escalation. CBID is not being asked to own NLP, "
        "statistical analysis, or humanitarian evaluation.",
        size=11,
    )
    add_para(
        doc,
        "Relevant team precedents: Elnakib et al., refugee integration into national health systems "
        "(Conflict and Health, 2024); Surmeli et al., digital reminders for antenatal care in a refugee "
        "population (JMIR preprint, 2026; NCT05094518); Shafieian, Abolfathi, and Yazdi, Caredesign "
        "(Annals of Biomedical Engineering, 2026).",
        size=11,
    )

    add_heading(doc, "Funding", 2)
    add_para(
        doc,
        "The working budget will include CBID faculty effort, the appropriate student-support mechanism, "
        "and allowable prototyping or facility costs if the collaboration is confirmed. We have deliberately "
        "not set the numbers; they should reflect CBID's mechanism and JHU requirements.",
        size=11,
    )

    add_heading(doc, "Questions we need your read on", 2)
    add_numbered(
        doc,
        [
            "Role: named co-investigator or faculty advisor?",
            "Student mechanism: MSE capstone team, undergraduate design team, research assistants, or a fellow? How many students, and which disciplines?",
            "Calendar: can one design cycle fit your June-to-May or academic-year cycle within our Months 13 to 24 window?",
            "Budget: what faculty effort, student support, prototyping and facility costs, and any CBID sponsorship or cost-recovery terms should we carry in the budget?",
            "Who else at CBID should join a scoping conversation?",
        ],
    )

    add_heading(doc, "Timing", 2)
    add_para(
        doc,
        "Submission closes July 15. CBID participation is described conditionally in the proposal until "
        "you confirm, so a 30-minute call this week is enough to name the role accurately.",
        size=11,
    )
    add_para(doc, "Contact: Shatha Elnakib (selnaki1@jhu.edu)  |  Aral Surmeli", bold=True, size=11)

    out = OUT_DIR / "lsri-yazdi-cbid-one-pager-2026-07-10.docx"
    doc.save(out)
    return out


def build_full_application() -> Path:
    doc = setup_doc(0.5)

    add_title(
        doc,
        "AI-Supported Digital Health Navigation and Conversational Analytics "
        "for Refugee Health Service Utilization",
    )
    add_subtitle(
        doc,
        "Lead PI: Shatha Elnakib, Center for Humanitarian Health, "
        "Johns Hopkins Bloomberg School of Public Health",
    )
    doc.add_paragraph()

    add_heading(doc, "Hypothesis and Specific Aims", 1)
    add_label_para(
        doc,
        "Hypothesis. ",
        "Routinely generated, de-identified conversational data from an operating digital health service "
        "can identify modifiable barriers to health service utilization among refugee and returnee "
        "populations, and a navigation intervention designed around one validated barrier can measurably "
        "improve navigation feasibility and service uptake.",
    )
    add_label_para(
        doc,
        "Aim 1. Establish and validate a governed conversational-and-utilization evidence pipeline. ",
        "We will audit and characterize the de-identified operational corpus (population, languages, volume, "
        "linkage to utilization indicators), implement governance and de-identification under JHU IRB, "
        "develop a barrier taxonomy using NLP and thematic coding, and evaluate lay-language health "
        "terminology in low-resource refugee-community languages where the corpus permits, validated against "
        "structured survey and utilization measures with prespecified convergence and divergence analyses. "
        "Outputs: governed analytic dataset; validated coding framework; terminology/accessibility brief where "
        "applicable; methods manuscript.",
    )
    add_label_para(
        doc,
        "Aim 2. Identify and prioritize modifiable barriers to public health service utilization. ",
        "Applying the Aim 1 pipeline, we will estimate barrier prevalence and associations with utilization "
        "patterns, triangulate conversational signals with structured measures, and select one priority "
        "barrier against prespecified criteria (prevalence, modifiability, equity relevance). Outputs: "
        "barrier evidence profile; cycle 1 learning memo; selected design target.",
    )
    add_label_para(
        doc,
        "Aim 3. Design and feasibility-test one AI-supported navigation intervention targeting the priority barrier. ",
        "We will translate Aim 2 findings into design requirements and, through a proposed CBID work package, "
        "produce one bounded navigation prototype integrated with the existing platform, then assess feasibility, "
        "acceptability, usability, and change in utilization indicators using rapid-cycle implementation evaluation. "
        "Outputs: prototype and design documentation; implementation findings; effect estimates to power follow-on studies.",
    )

    add_heading(doc, "1. Transformative Potential and Impact", 1)
    add_label_para(
        doc,
        "Significance. ",
        "Syrian refugee women in Turkiye retain formal access to public health services, including Ministry of "
        "Health migrant health centers, yet documented barriers to use persist, including language, navigation "
        "complexity, fragmented records, and mobility; clinical findings are not uniform across settings, and "
        "persistent maternal and newborn health inequalities remain across the evidence base.[1-6] Displacement "
        "compounds these barriers because health systems assume geographically stable users, while refugee and "
        "returnee families move across cities, providers, and documentation systems.[7] Prior digital maternal-health "
        "interventions in this and comparable populations, including the partner platform's earlier smartphone "
        "application for Syrian refugees, support the feasibility and acceptability of reminder and navigation support, "
        "though effects are not uniform across studies or components.[8-11] This study addresses a translational gap: "
        "humanitarian programs now generate large volumes of routine service conversations, but rigorous methods to link "
        "those data to validated barrier taxonomies and utilization outcomes in displacement contexts remain thin.[12,13]",
    )
    add_label_para(
        doc,
        "Innovation. ",
        "We will close a 24-month loop from ethically governed operational conversations to a validated barrier "
        "taxonomy to one deployed navigation intervention tested for feasibility and utilization change. The high-risk "
        "element is responsible use of sensitive, multilingual conversational data in populations for whom errors, privacy "
        "failures, and language or cultural bias carry disproportionate consequences.[14-17] Mainstream health-AI work "
        "rarely links consumer chatbot intent classification to subsequent care-seeking in humanitarian maternal and "
        "reproductive health populations.[12] We will test whether conversational signals add complementary barrier "
        "evidence beyond structured measures, whether lay-language terminology in underrepresented refugee-community "
        "languages (including dialectal Arabic where present in the corpus) aligns with clinical navigation constructs, "
        "and whether one barrier-informed navigation feature is feasible and associated with service use.",
    )
    add_label_para(
        doc,
        "Impact. ",
        "CHH will lead research design, ethics, analysis, and evaluation. HERA Digital Health will provide bounded "
        "de-identified data access and platform integration within the 10% external-subaward cap. This award funds "
        "catalytic science and implementation learning, not large-scale field implementation. Successful completion "
        "yields IRB-ready methods, an independent faculty-led research line in AI and humanitarian health evidence at CHH, "
        "a funded DrPH dissertation thread, supervised translational design experience for an anticipated CBID student team, "
        "peer-reviewable design and methods outputs, and a defensible evidence base for DIV, NIH, and foundation follow-on.",
    )

    add_heading(doc, "2. Strategic Alignment", 1)
    add_para(
        doc,
        "This proposal meets the Life Sciences Research Initiative mandate for bold, high-risk, high-reward science "
        "integrating data science and AI with population-health translation. AI is central to analytic methods and the "
        "navigation intervention; life-science relevance runs through maternal and reproductive health access among "
        "displaced populations; student investment is explicit through DrPH trainee support and a proposed CBID design team. "
        "The project advances CHH's capacity to generate rigorous evidence on emerging technology in humanitarian response.",
    )

    add_heading(doc, "3. Strength of Approach and Feasibility", 1)

    add_heading(doc, "Aim 1. Governed conversational-and-utilization evidence pipeline", 2)
    add_label_para(doc, "Rationale. ", "Aim 1 converts the operational corpus from a feasibility unknown into a validated analytic substrate. Without prespecified audit, governance, and validation against structured measures, later barrier prioritization and intervention design lack defensible inputs.")
    add_label_para(doc, "Design. ", "Mixed-methods corpus characterization and validation study over Months 1-12, nested within the 24-month award.")
    add_label_para(doc, "Data and measures. ", "De-identified conversational health data from HERA humanitarian operations; structured survey and utilization records where linkage is permitted. Prespecified audit documents population, languages, date range, usable volume, completeness, de-identification pathway, utilization linkage, permissions, and fitness for analysis.")
    add_table(
        doc,
        ["Outcome domain", "Candidate measures", "Role"],
        [
            ["Utilization (primary candidates)", "Antenatal visit counts; migrant health center registration and visit frequency; referral completion; sustained service engagement", "Primary where linkage permits; hierarchy finalized post-audit"],
            ["Conversational analytics", "Barrier taxonomy codes; thematic clusters; lay-language terminology fidelity in low-resource languages", "Aim 1 validation outputs"],
            ["Implementation (Aim 3)", "Feasibility; acceptability (AIM/IAM/FIM-class); usability (SUS); reach; fidelity", "Secondary in Aim 3"],
        ],
    )
    add_label_para(doc, "Analysis. ", "NLP and thematic coding of health-seeking narratives; pattern detection for entitlement confusion, facility access, and trust barriers; prespecified convergence and divergence analyses comparing conversational and structured signals; lay-language terminology review against clinician-reviewed navigation constructs where Arabic or other low-resource refugee-community languages are present in the corpus.")
    add_label_para(doc, "Pitfalls and alternatives. ", "Conversational data may add no useful information beyond structured measures or may encode platform-selection and documentation biases; we treat these as testable null hypotheses in Aim 1 rather than assumptions.")

    add_heading(doc, "Aim 2. Barrier identification and prioritization", 2)
    add_label_para(doc, "Rationale. ", "Utilization improvement requires targeting modifiable barriers with equity relevance, not generic access narratives.")
    add_label_para(doc, "Design. ", "Cross-sectional and longitudinal association analyses using the Aim 1 dataset, followed by structured barrier prioritization against prespecified criteria (prevalence, modifiability, equity relevance).")
    add_label_para(doc, "Analysis. ", "Estimate barrier prevalence and associations with utilization patterns; triangulate conversational and structured evidence; produce cycle 1 learning memo and one CHH-approved design target for Aim 3.")

    add_heading(doc, "Aim 3. Navigation intervention design and feasibility testing", 2)
    add_label_para(doc, "Rationale. ", "The translational test is whether research-derived barrier evidence can become a usable, safe navigation response on an operating platform within the award period.")
    add_label_para(doc, "Design. ", "Rapid-cycle implementation evaluation over Months 13-24: design requirements, bounded prototype, supervised usability and feasibility testing, linkage to utilization indicators where feasible. We do not propose a large RCT; the program does not require one.")
    add_label_para(doc, "CBID translational work package. ", "The proposed CBID collaboration applies CBID's established partner-anchored student design model, recently demonstrated in the 2024 Ayu AI collaboration with the NGO telemedicine platform Intelehealth, and follows CBID's published needs-finding methodology for healthcare service innovation in low- and middle-income countries.[18]")
    add_numbered(
        doc,
        [
            "Synthesize Aim 2 findings into user, accessibility, and system requirements.",
            "Map the patient and navigation journey around the priority barrier.",
            "Generate and assess intervention concepts with CHH, HERA, and the DrPH trainee.",
            "Prototype one bounded HERA-compatible navigation feature or workflow on the existing WhatsApp-based platform.",
            "Conduct supervised usability and feasibility testing under the approved ethics and partner protocol.",
            "Deliver design documentation, implementation recommendations, and handoff for follow-on funding.",
        ],
    )
    add_table(
        doc,
        ["Phase", "Activity", "Output"],
        [
            ["Year 1, Q1-Q2", "IRB; de-identification; corpus audit; baseline utilization metrics", "Approved protocol; analytic dataset"],
            ["Year 1, Q3-Q4", "Aim 1 validation; Aim 2 barrier prioritization; design brief", "Coding framework; cycle 1 learning memo"],
            ["Year 2, Q1-Q2", "Aim 3 concept generation and bounded prototype", "Prototype; design documentation"],
            ["Year 2, Q3-Q4", "Supervised usability/feasibility; utilization linkage", "Implementation findings; effect estimates for power calculations"],
        ],
    )
    add_label_para(doc, "Feasibility. ", "JHU hosts analysis, faculty effort, trainee support, and compute. HERA provides data access and technical partnership within the 10% subaward cap. Prior quasi-experimental work on antenatal visit attendance in Istanbul migrant health centers demonstrates adjacent evaluation capacity, though that study used smartphone reminders rather than conversational AI and supports plausibility only.[10]")

    add_heading(doc, "4. Team Capability and Environment", 1)
    add_label_para(doc, "Lead PI. ", "Shatha Elnakib's scholarship spans the scientific bridge this proposal requires. Her mixed-methods work examines how refugees are integrated into national health systems and why policy commitments do or do not translate into sustained access.[19] Her humanitarian sexual, reproductive, maternal, and newborn health research includes evaluation of service delivery during conflict.[20] She has also contributed to validation of a mobile-health mortality-risk model using national survey data.[21] Elnakib will be accountable for study design, IRB and ethics, the qualitative codebook and validation framework, mixed-methods integration, selection of the Aim 3 design target, and publication.")
    add_label_para(doc, "Trainee. ", "Aral Surmeli is a DrPH candidate advised by Paul Spiegel and the physician-founder of HERA Digital Health. He is first author of the current JMIR preprint and registered principal investigator for NCT05094518, which evaluated antenatal-care reminders among Syrian refugee women at two Istanbul migrant health centers.[10] Under Elnakib's supervision, Surmeli will lead operational data-provenance mapping, linkage-feasibility documentation, implementation protocol and fidelity assessment, and rapid-cycle learning. His earlier publications on HERA's refugee-health model and end-user acceptability provide additional population and implementation context.[8,9]")
    add_label_para(doc, "Complementarity and development. ", "Elnakib defines the scientific and ethical standard and determines which finding is sufficiently validated to translate. Surmeli makes the operational data and implementation pathway auditable. Subject to confirmation, Yazdi will provide design-methodology oversight and student supervision while a CBID team converts one approved finding into a testable navigation workflow. HERA supplies bounded data and integration support within the 10% external-partner cap.")
    add_label_para(doc, "Environment. ", "JHU CHH provides humanitarian health evaluation rigor and scientific ownership. The proposed CBID collaboration adds bioengineering design methods. HERA provides the operational platform and de-identified data interface within the 10% cap without new field operations.")

    add_heading(doc, "5. Future Funding Potential", 1)
    add_para(doc, "This Individual Award is designed as a springboard. Follow-on pathways include DIV Fund stages, NIH and foundation implementation grants, and CHH-HERA co-authored methods and design papers.")

    add_heading(doc, "6. Appropriate Use of Institutional Funds and Budget Justification", 1)
    add_label_para(doc, "Requested amount: ", "$300,000 over 2 years (range $200k-$500k allowed; confirm with Shatha/JHURA).")
    add_table(
        doc,
        ["Line item", "Amount (placeholder)", "Justification"],
        [
            ["PI effort (Shatha)", "$45,000", "15% effort Years 1-2; study design, oversight, publication"],
            ["CBID faculty and student participation", "TBD", "Faculty effort, student mechanism, and facilities per Yazdi/Shatha/JHURA scoping"],
            ["Trainee support (Aral)", "$30,000", "Dissertation-related research activities"],
            ["Postdoc / research analyst (CHH)", "$90,000", "Thematic analysis, implementation learning cycles"],
            ["Compute / AI infrastructure", "$50,000", "NLP pipelines, secure storage, JHU-hosted compute"],
            ["Prototyping, usability, travel, dissemination", "TBD", "Prototype materials, supervised usability, partner coordination"],
            ["HERA partnership (subaward)", "$30,000", "10% cap at $300k: de-identification support, API access, technical deliverables; or in-kind"],
            ["Working request envelope", "$300,000", "Rebalance after CBID and JHURA scoping"],
        ],
    )

    add_heading(doc, "7. Demonstrated Need for Support", 1)
    add_para(doc, "This project is not yet competitive for traditional external mechanisms at full implementation scale. The Individual Award provides catalytic investment to establish IRB-ready conversational analytics methods in humanitarian settings, generate utilization-focused implementation evidence, and position the faculty team for external follow-on.")

    add_heading(doc, "References (beyond 4-page limit)", 1)
    refs = [
        "Davidson N, Hammarberg K, Romero L, Fisher J. Access to preventive sexual and reproductive health care for women from refugee-like backgrounds: a systematic review. BMC Public Health. 2022;22:403.",
        "Ozel S, Yaman S, Kansu-Celik H, Hancerliogullari N, Balci N, Engin-Ustun Y. Obstetric outcomes among Syrian refugees: a comparative study at a tertiary care maternity hospital in Turkey. Revista Brasileira de Ginecologia e Obstetricia. 2018;40(11):673-679.",
        "Sayili U, Ozgur C, Bulut Gazanfer O, Solmaz A. Comparison of clinical characteristics and pregnancy and neonatal outcomes between Turkish citizens and Syrian refugees with high-risk pregnancies. Journal of Immigrant and Minority Health. 2022;24(5):1177-1185.",
        "Hakimi S, et al. Maternal and newborn health inequality among Syrian refugees in Turkey: a systematic review and meta-analysis. International Journal for Equity in Health. 2025;24:160.",
        "Republic of Turkiye Ministry of Interior, Presidency of Migration Management. Rights and obligations under temporary protection. Accessed July 10, 2026.",
        "Republic of Turkiye Ministry of Health, General Directorate of Public Health. Migrant health centers. Accessed July 10, 2026.",
        "Abubakar I, Aldridge RW, Devakumar D, et al. The UCL-Lancet Commission on Migration and Health: the health of a world on the move. The Lancet. 2018;392(10164):2606-2654.",
        "Surmeli A, Narla NP, Shields AJ, Atun R. Leveraging mobile applications in humanitarian crisis to improve health: a case of Syrian women and children refugees in Turkey. Journal of Global Health Reports. 2020;4:e2020099.",
        "Meyer CL, Surmeli A, Hoeflin Hana C, Narla NP. Perceptions on a mobile health intervention to improve maternal child health for Syrian refugees in Turkey. Frontiers in Public Health. 2022;10:1025675.",
        "Surmeli A, Yasin Y, Narla NP, et al. Digital reminders for antenatal care in a refugee population. JMIR Preprints. 2026;97767. Under peer review; NCT05094518.",
        "Saleh S, El Arnaout N, Sabra N, et al. The effectiveness of an artificial intelligence-based gamified intervention for improving maternal health outcomes among refugees and underserved women in Lebanon. JMIR mHealth and uHealth. 2025;13:e65599.",
        "Costa-Gomes B, Tolmachev P, Taysom E, et al. Public use of a generalist LLM chatbot for health queries. Nature Health. 2026;1:689-696.",
        "Norton A, Tappis H. Sexual and reproductive health implementation research in humanitarian contexts: a scoping review. Reproductive Health. 2024;21:64.",
        "World Health Organization. The role of artificial intelligence in sexual and reproductive health and rights: technical brief. Geneva: WHO; 2024.",
        "Cirillo D, et al. Sex and gender differences and biases in artificial intelligence for biomedicine and healthcare. npj Digital Medicine. 2020;3:81.",
        "Zhu L, Mou W, Lai Y, Lin J, Luo P. Language and cultural bias in AI. Journal of Translational Medicine. 2024;22:319.",
        "Huo B, Boyle A, Marfo N, et al. Large language models for chatbot health advice studies: a systematic review. JAMA Network Open. 2025;8(2):e2457879.",
        "Shafieian M, Abolfathi N, Yazdi Y. Caredesign: a needs-finding framework for enhancing healthcare service innovation in low- and middle-income countries. Annals of Biomedical Engineering. 2026;54(2):613-620.",
        "Elnakib S, Jackson C, Lalani U, Shawar YR, Bennett S. How integration of refugees into national health systems became a global priority. Conflict and Health. 2024;18(Suppl 1):31.",
        "Tappis H, Elaraby S, Elnakib S, et al. Reproductive, maternal, newborn and child health service delivery during conflict in Yemen: a case study. Conflict and Health. 2020;14:30.",
        "Elnakib S, Vecino-Ortiz AI, Gibson DG, et al. A novel score for mHealth apps to predict and prevent mortality. Journal of Medical Internet Research. 2022;24(6):e36787.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {ref}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    out = OUT_DIR / "lsri-application-narrative-2026-07-10.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    y = build_yazdi_one_pager()
    a = build_full_application()
    print(f"Wrote: {y}")
    print(f"Wrote: {a}")
